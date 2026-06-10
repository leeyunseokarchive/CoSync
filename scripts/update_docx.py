from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import re

DOC_PATH = '/Users/t2025-m0051/Desktop/CoSync_온보딩진단_기획서.docx'

doc = Document(DOC_PATH)
paras = doc.paragraphs

# --- 1. 심화 진단 개수 설명 수정 (para 7) ---
for p in paras:
    if '심화 진단: 8문항 (1개 카테고리)' in p.text:
        for run in p.runs:
            run.text = ''
        p.runs[0].text = '• 심화 진단: 8문항 (2개 카테고리) — 선택, 1차 완료 후 진입 가능'
        break

# --- 2. para 117부터 끝까지 삭제 후 새 내용 삽입 ---
# 117: "3. 심화 진단"부터 끝까지 전부 교체

# 삭제할 단락 인덱스 찾기
start_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '3. 심화 진단 (선택, 8문항)':
        start_idx = i
        break

if start_idx is None:
    print("시작 단락을 못 찾음")
    exit(1)

# body element 참조
body = doc.element.body

# 삭제할 단락들 수집
to_remove = list(doc.paragraphs[start_idx:])
for p in to_remove:
    p._element.getparent().remove(p._element)

# --- 새 내용 추가 헬퍼 ---
def add_para(doc, text, bold=False, style='Normal'):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    return p

def add_divider(doc):
    add_para(doc, '─' * 60)

def add_question(doc, qnum, title, scenario, options, toxic):
    add_divider(doc)
    add_para(doc, f'{qnum}. {title}', bold=True)
    add_para(doc, '시나리오+질문:')
    add_para(doc, scenario)
    for opt in options:
        add_para(doc, opt)
    add_para(doc, f'※ Toxic pair: {toxic}')

# --- 3. 심화 진단 섹션 ---
add_para(doc, '3. 심화 진단 (선택, 8문항)', bold=True)
add_para(doc, '• 1차 진단(Q1~Q12) 완료 후 모달에서 "심화 진단 계속하기" 선택 시 진입')
add_para(doc, '• 완료 시 갭 리포트에 심화 히트맵(비전/가치관, 돈/보상) 추가 노출')
add_para(doc, '• 두 사람 모두 완료해야 심화 갭 리포트 활성화')
add_para(doc, '')

# 카테고리 4
add_para(doc, '카테고리 4. 비전/가치관', bold=True)
add_para(doc, '가중치: 1.5 | 측정 목적: 출구 전략·피벗 기준·갈등 처리 방식의 철학 충돌 파악')
add_para(doc, '')

add_question(doc, 'Q13', '회사 출구 전략',
    '창업을 시작한 지 2년이 지났고 제품이 시장에서 반응을 얻기 시작했습니다. 이 회사의 가장 이상적인 결말은 무엇이라고 생각하나요?',
    ['1. 빠르게 성장해 대기업에 인수된다 (M&A 엑싯)',
     '2. 상장(IPO)해서 더 큰 회사로 키운다',
     '3. 외부 투자 없이 수익성 있는 독립 회사로 오래 운영한다',
     '4. 아직 정해진 생각은 없다'],
    '1 ↔ 3 (M&A 엑싯 vs 독립 운영 — 출구 전략 근본 충돌)')

add_question(doc, 'Q14', '피벗/중단 기준',
    '팀이 목표한 성과를 지속적으로 내지 못하고 있습니다. 중단 또는 방향 전환을 해야 한다고 판단하는 기준은 무엇인가요?',
    ['1. 자금이 부족해 더 이상 운영이 어려울 때',
     '2. 일정 기간 동안 시장 반응이 거의 없을 때',
     '3. 핵심 팀원이 이탈하거나 지속적으로 흔들릴 때',
     '4. 공동창업자 간 합의가 되지 않을 때'],
    '1 ↔ 2 (자금 소진까지 버팀 vs 시장 신호로 먼저 판단)')

add_question(doc, 'Q15', '갈등 해소 방식',
    '공동창업자 간 갈등이 생겼을 때 어떻게 해결하는 게 맞다고 생각하나요?',
    ['1. 당사자끼리 즉시 직접 대화로 해결한다',
     '2. 정해진 기준이나 룰에 따라 처리한다',
     '3. 냉각 기간을 두고 나서 이야기한다',
     '4. 외부 멘토나 제3자의 도움을 받는다'],
    '1 ↔ 3 (즉시 직접 대화 vs 냉각 기간 후 대화 — 갈등 처리 스타일 충돌)')

add_question(doc, 'Q16', '절대 용납 못하는 것',
    '공동창업자에게 절대 용납할 수 없는 것은 무엇인가요?',
    ['1. 결정을 미루거나 느리게 움직이는 것',
     '2. 말한 것을 지키지 않는 것',
     '3. 결과 없이 이유만 대는 것',
     '4. 방향이 달라지고 있는데 맞추려 하지 않는 것'],
    '1 ↔ 4 (속도를 못 참는 것 vs 방향 불일치를 못 참는 것)')

add_para(doc, '')

# 카테고리 5
add_para(doc, '카테고리 5. 돈/보상', bold=True)
add_para(doc, '가중치: 2.0 | 측정 목적: 급여·지분·수익 배분 철학의 충돌 파악')
add_para(doc, '')

add_question(doc, 'Q17', '급여 구조',
    '창업 초기 자금이 넉넉하지 않습니다. 공동창업자 간 급여는 어떻게 정하는 게 맞다고 생각하나요?',
    ['1. 역할과 기여도에 따라 처음부터 차등 지급한다',
     '2. 초기엔 동일하게 맞추고 이후 성과에 따라 조정한다',
     '3. 회사가 안정될 때까지 최소 수준으로 맞춘다',
     '4. 각자 시장 기준 연봉에 맞게 책정한다'],
    '1 ↔ 2 (역할 차등 지급 vs 동일 분배 — 보상 철학 충돌)')

add_question(doc, 'Q18', '지분 구조 철학',
    '공동창업자 지분 구조에 대해 어떻게 생각하나요?',
    ['1. 투자 유치를 위해 시장 관행에 맞는 구조를 유지하는 게 맞다',
     '2. 실제 기여도와 시간 투입이 달라지면 지분도 조정해야 한다',
     '3. 처음 합의한 지분은 어떤 상황에서도 계약대로 이행해야 한다',
     '4. 지분은 고정하되 스톡옵션이나 급여로 기여도 차이를 메운다'],
    '1 ↔ 2 (관행 고수 vs 기여 기반 유동 조정 — 지분 철학 충돌)')

add_question(doc, 'Q19', '수익 배분 우선순위',
    '회사에 의미 있는 수익이 발생하기 시작했습니다. 이 수익을 어떻게 처리하는 게 맞다고 생각하나요?',
    ['1. 전액 사업에 재투자한다 — 지금은 성장이 먼저다',
     '2. 일부는 재투자, 일부는 공동창업자 보상으로 배분한다',
     '3. 급여 인상이나 인센티브를 먼저 챙긴다',
     '4. 적립해두고 팀이 함께 결정할 때 쓴다'],
    '1 ↔ 3 (전액 재투자 vs 보상 먼저 — 수익 철학 충돌)')

add_question(doc, 'Q20', '성장 전략',
    '회사를 성장시키는 방식에 대해 어떻게 생각하나요?',
    ['1. 외부 투자를 받아 빠르게 성장한다 — 지분 희석은 감수한다',
     '2. 수익으로 버티면서 최대한 지분을 지킨다',
     '3. 필요한 시점에 선택적으로 투자를 받는다',
     '4. 정부 지원금이나 대출 등 비희석 자금을 먼저 찾는다'],
    '1 ↔ 2 (희석 감수 성장 vs 지분 보호 생존 — 성장 전략 충돌)')

add_para(doc, '')

# --- 4. 카테고리 가중치 구조 ---
add_para(doc, '4. 카테고리 가중치 구조', bold=True)
add_para(doc, '• 권한 & 실행 (Q1~Q4): weight 1.0')
add_para(doc, '• 책임 (Q5~Q8): weight 1.5')
add_para(doc, '• 종료 (Q9~Q12): weight 2.0')
add_para(doc, '• 비전/가치관 (Q13~Q16): weight 1.5')
add_para(doc, '• 돈/보상 (Q17~Q20): weight 2.0')
add_para(doc, '')
add_para(doc, '갭 점수 기준 (최대 가능 점수 102점):')
add_para(doc, '• CRITICAL: 32점 이상')
add_para(doc, '• HIGH: 18~31점')
add_para(doc, '• MID: 7~17점')
add_para(doc, '• LOW: 6점 이하')
add_para(doc, '')

# --- 5. UX 플로우 ---
add_para(doc, '5. UX 플로우', bold=True)
add_para(doc, '1. 온보딩 진입 → 1차 진단 시작 (Q1~Q12, 약 5분)')
add_para(doc, '2. 카테고리 1~3 순서로 진행 (진행바 표시)')
add_para(doc, '3. Q12 완료 → 심화 진단 선택 모달 표시')
add_para(doc, '   - [심화 진단 계속하기 (Q13~Q20) →]')
add_para(doc, '   - 팀 없는 경우: [팀 만들고 공동창업자 초대하기]')
add_para(doc, '   - 팀 있는 경우: [결과 바로 확인하기]')
add_para(doc, '4. 심화 선택 시 카테고리 4~5 진행 → 프로필 완료 → 워크스페이스')
add_para(doc, '5. 갭 리포트: 1차 히트맵(Q1~Q12) 기본 표시')
add_para(doc, '6. 두 사람 모두 심화 완료 시 오른쪽에 심화 히트맵(Q13~Q20) 추가 노출')
add_para(doc, '')

# --- 6. 미결 사항 ---
add_para(doc, '6. 미결 사항', bold=True)
add_para(doc, '• 심화 진단 진입 이후 중단 시 재진입 UX 처리')
add_para(doc, '• 심화 갭 리포트 AI 분석 텍스트 추가')

doc.save(DOC_PATH)
print("저장 완료:", DOC_PATH)
